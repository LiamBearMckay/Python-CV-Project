from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture(
    'pexels-photo-771742.jpeg', 
    width=Inches(2.0))

# name, phone and email details
name = input('what is your name? ')
speak('hello' + name + 'How are you today?')

speak('what is your phone number?')
phone = input('what is your phone number? ')
email = input('what is your email? ')

document.add_paragraph(
   name +  '  | ' + phone + ' | ' + email)

# about me 
document.add_heading('About me')
about_me = input('tell me about yourself? ')
document.add_paragraph(about_me)


# work experience 
document.add_heading('work experience')
p = document.add_paragraph()

company = input('enter company ')
from_date = input('from date ')
to_date = input('to date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'describe your experience at ' + company)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'do you have any more experience? yes or no? ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('enter company ')
        from_date = input('from date ')
        to_date = input('to date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
        'describe your experience at ' + company)
        p.add_run(experience_details)
    else:
        break

#footer
section = document.sections[0]
footer = section.footer 
p = footer.paragraphs[0]
p.text = 'CV generated in python'

document.save('cv.docx')